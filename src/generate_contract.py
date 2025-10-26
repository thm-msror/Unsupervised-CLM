import os
import json
import re
import time
from dotenv import load_dotenv
import google.generativeai as genai

from doc_reader import read_docu
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches


def generate_contract(contract_data: dict, markdown_template: str, old_contract: str = "") -> str:
    """
    Generates a full professional contract in Markdown format using Gemini 2.5 Flash.
    """
    load_dotenv()
    api_key = os.getenv("GEMINI_API_KEY")
    genai.configure(api_key=api_key)

    if old_contract:
        old_contract_prompt = f"""
        ### 📄 Old contract:
        ### CRITICAL INSTRUCTION: You will use this old contract as a base for the new contract. You will try to use as much from the old as possible. You will only change sentences if the old is not compliant with the **{contract_data["Industry"]}** industry standards and the **{contract_data["Jurisdiction"]}** 
        {old_contract}

        ---
        """
    else: 
        old_contract_prompt = ""
    

    prompt = f"""
        You are a professional contract automation AI.

        Use the JSON data and Markdown structure below to generate a **complete, legally sound, and jurisdiction-specific contract** in Markdown format. 
        Fill all placeholders naturally and professionally — every section must align with the provided context.

        ---

        ### Contract Data (JSON)
        {json.dumps(contract_data, indent=2)}

        ---

        ### 📄 Markdown Template
        {markdown_template}

        ---

        {old_contract_prompt}

        ### CRITICAL INSTRUCTIONS — FOLLOW STRICTLY
        - You MUST adapt all clauses, tone, and terminology to **the specified Industry: "{contract_data["Industry"]}"**.  
        Example: a tech contract should include IP, confidentiality, and data protection clauses; a construction contract should include project scope, timelines, and site liability.
        - You MUST adapt all legal phrasing, governing law, and enforcement clauses to **the specified Jurisdiction: "{contract_data["Jurisdiction"]}"**.  
        Example: use region-appropriate terminology, legal references, and dispute resolution norms.
        - Do NOT include or reference any other jurisdiction, law, or country outside "{contract_data["Jurisdiction"]}".
        - Do NOT generalize clauses. Always ground them in the provided Industry and Jurisdiction context.
        - Replace all placeholders like {{PartyA.LegalName}} or {{EffectiveDate}} with actual JSON data.
        - Maintain Markdown structure and headings (e.g., `#`, `##`, `###`) exactly as given in the template.
        - The contract must be written entirely in **{contract_data["Language"]}**.
        - Output **only** the completed Markdown contract — no explanations, no notes.

        ---

        ### GOAL
        Produce a polished, realistic, and enforceable contract fully compliant with the **{contract_data["Industry"]}** industry standards and the **{contract_data["Jurisdiction"]}** legal system.
        """
    

    model = genai.GenerativeModel("gemini-2.5-flash")
    response = model.generate_content(prompt)
    return response.text.strip()

def markdown_to_docx(text: str, output_path: str):
    """
    Converts Markdown-style contract text to a properly formatted, professional Word document.
    Supports headings (#), bold (**), italics (*), and normal paragraphs.
    """
    doc = Document()

    # Global style
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    lines = text.splitlines()
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        # --- HEADINGS WITH FORMATTING ---
        if stripped.startswith("### "):
            heading = stripped.replace("### ", "").strip()
            p = doc.add_paragraph()
            run = p.add_run(heading)
            run.bold = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)   # Heading 3: 12–13 pt
            run.font.color.rgb = RGBColor(0, 0, 0)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(4)

        elif stripped.startswith("## "):
            heading = stripped.replace("## ", "").strip()
            p = doc.add_paragraph()
            run = p.add_run(heading.upper())  # optional all caps
            run.bold = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(14)   # Heading 2: 14 pt
            run.font.color.rgb = RGBColor(0, 0, 0)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(6)

        elif stripped.startswith("# "):
            heading = stripped.replace("# ", "").strip()
            p = doc.add_paragraph()
            run = p.add_run(heading.upper())
            run.bold = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(16)   # Heading 1: 16–18 pt
            run.font.color.rgb = RGBColor(0, 0, 0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.space_before = Pt(30)


        else:
            # Convert bold and italics manually
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            tokens = re.split(r"(\*\*.*?\*\*|\*.*?\*)", stripped)
            for token in tokens:
                run = p.add_run()
                if token.startswith("**") and token.endswith("**"):
                    run.text = token[2:-2]
                    run.bold = True
                elif token.startswith("*") and token.endswith("*"):
                    run.text = token[1:-1]
                    run.italic = True
                else:
                    run.text = token

        p.paragraph_format.space_after = Pt(6)
        

    # Save as .docx
    doc.save(output_path)
    print(f"✅ Styled contract saved at: {output_path}")
    
def generate_contract_doc(contract_data, template_path, output_file_path, md_path=False, old_contract_path = False):
    start_time = time.time()

    # Read template_path
    with open(template_path, "r", encoding="utf-8") as f:
        markdown_template = f.read()

    if old_contract_path:
        old_contract= read_docu(old_contract_path)["full_text"]

    else:
        old_contract = ""


    #============================================
    # Generating contract
    final_contract = generate_contract(contract_data, markdown_template, old_contract)
    #============================================

    # Make output_file_path if missing
    output_folder = os.path.dirname(output_file_path)
    os.makedirs(output_folder, exist_ok=True)

    # Save as MD
    if md_path:
        with open(md_path, "w", encoding="utf-8") as f:
            f.write(final_contract)
        print(f"✅ Markdown contract saved as: {md_path}")
    
    # Save as Word Document
    markdown_to_docx(final_contract, output_file_path)
    print(f"✅ Word contract saved as: {output_file_path}")

    print(f"⏱️ Contract generation completed in {time.time() - start_time:.2f} seconds.")

def test():
    contract_input_1 = {
        "ContractTemplate": "Services Agreement",
        "Industry": "Technology",
        "Jurisdiction": "Qatar",
        "Language": "English",
        "Start_date": "2025-11-01",
        "End_date": "2026-11-01",
        "PartyA": {
            "LegalName": "QTech Solutions W.L.L.",
            "Role": "Client",
            "Address": "Office 304, Doha Tower, West Bay, Doha, Qatar"
        },
        "PartyB": {
            "LegalName": "CodeBridge Technologies Ltd.",
            "Role": "Contractor",
            "Address": "45 Innovation Street, London, United Kingdom"
        },
        "Context": "Software development and maintenance of a web-based platform for QTech’s internal operations."
    }
   
    generate_contract_doc(
        contract_data = contract_input_1, 
        template_path = r"prompts\starter_contract.md", 
        output_file_path = r"data\generated_contracts\contract_output.docx", 
        old_contract_path=r"data\english\FinAmFamMut2016RestructExhA-19.docx"
    )

test()